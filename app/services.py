from __future__ import annotations

import logging
import uuid
from pathlib import Path
from typing import Any

from sqlalchemy import select, or_, update as sa_update
from sqlalchemy.orm import Session, joinedload

from app.models import (
    KBChunk, KBDocument, KBFolder, KnowledgeBase,
    User, SystemSetting,
)
from app.schemas import KBFolderTreeNode
from app.core.security import hash_password, verify_password
from app.vector_store import get_vector_store, DEFAULT_RAG_CONFIG

logger = logging.getLogger(__name__)

# Where uploaded files live on disk
UPLOAD_DIR = Path(__file__).resolve().parents[1] / "uploads"
UPLOAD_DIR.mkdir(exist_ok=True)

# Legacy ChromaDB path — kept for backward compatibility migrations.
# New code should use get_vector_store() from app.vector_store.
CHROMA_DIR = Path(__file__).resolve().parents[1] / "chroma_db"
CHROMA_DIR.mkdir(exist_ok=True)


# ============================================================
# File-type detection
# ============================================================

def detect_file_type(filepath: str) -> str:
    ext = Path(filepath).suffix.lower()
    mapping = {
        ".pdf": "pdf", ".docx": "docx", ".txt": "txt", ".md": "markdown",
        ".csv": "csv", ".json": "json", ".py": "code", ".js": "code",
        ".ts": "code", ".jsx": "code", ".tsx": "code", ".html": "code",
        ".css": "code", ".java": "code", ".go": "code", ".rs": "code",
        ".yaml": "code", ".yml": "code", ".toml": "code", ".xml": "code",
        ".sh": "code", ".bash": "code",
    }
    return mapping.get(ext, "unknown")


# ============================================================
# Text extraction helpers
# ============================================================

def extract_text_from_file(filepath: str, file_type: str) -> tuple[str, int]:
    if file_type == "pdf":
        return _extract_pdf(filepath)
    if file_type == "docx":
        return _extract_docx(filepath)
    if file_type in ("markdown", "txt", "csv", "json", "code"):
        return _extract_text(filepath)
    return ("", 0)


def _extract_pdf(filepath: str) -> tuple[str, int]:
    try:
        from pdfminer.high_level import extract_text
        text = extract_text(filepath)
        return (text or "", 1)
    except Exception as exc:
        logger.warning("pdfminer failed: %s", exc)
        try:
            import pypdf
            reader = pypdf.PdfReader(filepath)
            pages = [page.extract_text() or "" for page in reader.pages]
            return ("\n\n".join(pages), len(reader.pages))
        except Exception as exc2:
            logger.warning("pypdf also failed: %s", exc2)
            return ("", 0)


def _extract_docx(filepath: str) -> tuple[str, int]:
    try:
        from docx import Document
        doc = Document(filepath)
        text = "\n".join(p.text for p in doc.paragraphs if p.text)
        return (text, 1)
    except Exception as exc:
        logger.warning("python-docx failed: %s", exc)
        return ("", 0)


def _extract_text(filepath: str) -> tuple[str, int]:
    try:
        text = Path(filepath).read_text(encoding="utf-8", errors="replace")
        return (text, 1)
    except Exception as exc:
        logger.warning("text read failed: %s", exc)
        return ("", 0)


# ============================================================
# Chunking
# ============================================================

def chunk_text(text: str, chunk_size: int = 500, chunk_overlap: int = 50) -> list[dict]:
    # Lazy import: only load when chunking is actually needed
    from langchain_text_splitters import RecursiveCharacterTextSplitter

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", ". ", "。", " ", ""],
        length_function=len,
    )
    docs = splitter.create_documents([text], metadatas=[{}])
    return [{"content": d.page_content, "metadata": d.metadata} for d in docs]


# ============================================================
# Embeddings
# ============================================================

def get_embeddings(model_name: str = "text-embedding-3-small"):
    # Lazy import: only load when embeddings are actually needed
    from langchain_openai import OpenAIEmbeddings

    from app.settings import get_settings
    settings = get_settings()
    return OpenAIEmbeddings(
        model=model_name,
        api_key=settings.openai_api_key,
        base_url=settings.openai_base_url,
    )


# ============================================================
# Knowledge Base Service
# ============================================================

class KnowledgeBaseService:
    """All knowledge-base business logic."""

    @staticmethod
    def create_kb(db: Session, user_id: int, name: str, description: str = "",
                  embedding_model: str = "text-embedding-3-small",
                  chunk_size: int = 500, chunk_overlap: int = 50,
                  enabled: bool = True) -> KnowledgeBase:
        kb = KnowledgeBase(
            user_id=user_id, name=name, description=description,
            embedding_model=embedding_model, chunk_size=chunk_size,
            chunk_overlap=chunk_overlap, enabled=enabled,
            rag_config=DEFAULT_RAG_CONFIG.copy(),
        )
        db.add(kb)
        db.commit()
        db.refresh(kb)
        return kb

    @staticmethod
    def get_kb(db: Session, kb_id: int, user_id: int) -> KnowledgeBase | None:
        return db.scalar(
            select(KnowledgeBase).where(KnowledgeBase.id == kb_id, KnowledgeBase.user_id == user_id)
        )

    @staticmethod
    def list_kbs(db: Session, user_id: int) -> list[KnowledgeBase]:
        return list(db.scalars(
            select(KnowledgeBase).where(KnowledgeBase.user_id == user_id).order_by(KnowledgeBase.created_at)
        ))

    @staticmethod
    def update_kb(db: Session, kb: KnowledgeBase, **kwargs) -> KnowledgeBase:
        for k, v in kwargs.items():
            if v is not None and hasattr(kb, k):
                setattr(kb, k, v)
        db.commit()
        db.refresh(kb)
        return kb

    @staticmethod
    def delete_kb(db: Session, kb: KnowledgeBase) -> None:
        try:
            vs = get_vector_store()
            vs.delete_collection(f"kb_{kb.id}")
        except Exception as exc:
            logger.warning("Failed to delete vector collection for kb %s: %s", kb.id, exc)
        db.delete(kb)
        db.commit()

    @staticmethod
    def create_folder(db: Session, kb_id: int, name: str, description: str = "",
                      parent_id: int | None = None) -> KBFolder:
        folder = KBFolder(kb_id=kb_id, name=name, description=description, parent_id=parent_id)
        db.add(folder)
        db.commit()
        db.refresh(folder)
        return folder

    @staticmethod
    def delete_folder(db: Session, folder: KBFolder) -> None:
        db.delete(folder)
        db.commit()

    @staticmethod
    def get_folder_tree(db: Session, kb_id: int) -> list[KBFolderTreeNode]:
        from sqlalchemy.orm import undefer
        folders = db.scalars(
            select(KBFolder)
            .options(joinedload(KBFolder.children))
            .where(KBFolder.kb_id == kb_id)
        ).unique().all()

        def build_node(folder: KBFolder) -> KBFolderTreeNode:
            doc_count = len(folder.documents)
            return KBFolderTreeNode(
                id=folder.id, name=folder.name, description=folder.description,
                children=[build_node(child) for child in folder.children],
                document_count=doc_count,
            )

        return [build_node(f) for f in folders]

    @staticmethod
    def get_folder_path(db: Session, folder: KBFolder) -> str:
        parts: list[str] = []
        current = folder
        while current:
            parts.append(current.name)
            parent = db.get(KBFolder, current.parent_id) if current.parent_id else None
            current = parent
        parts.reverse()
        return " / ".join(parts)

    @staticmethod
    def upload_document(db: Session, kb_id: int, user_id: int, folder_id: int | None,
                        file_bytes: bytes, original_filename: str) -> KBDocument:
        file_type = detect_file_type(original_filename)
        if file_type == "unknown":
            raise ValueError(f"Unsupported file type: {original_filename}")
        safe_name = uuid.uuid4().hex[:12] + "_" + original_filename
        dest = UPLOAD_DIR / safe_name
        dest.write_bytes(file_bytes)
        doc = KBDocument(
            kb_id=kb_id, folder_id=folder_id, user_id=user_id,
            original_filename=original_filename, storage_path=str(dest),
            file_type=file_type, file_size=len(file_bytes), status="pending",
        )
        db.add(doc)
        db.flush()
        return doc

    @staticmethod
    def list_documents(db: Session, kb_id: int, folder_id: int | None = None) -> list[KBDocument]:
        q = select(KBDocument).where(KBDocument.kb_id == kb_id)
        if folder_id is not None:
            q = q.where(KBDocument.folder_id == folder_id)
        return list(db.scalars(q.order_by(KBDocument.created_at.desc())))

    @staticmethod
    def delete_document(db: Session, doc: KBDocument) -> None:
        try:
            Path(doc.storage_path).unlink(missing_ok=True)
        except Exception:
            pass
        chunks = db.scalars(select(KBChunk).where(KBChunk.document_id == doc.id)).all()
        if chunks:
            kb = doc.kb
            try:
                vs = get_vector_store()
                vs.delete(f"kb_{kb.id}", ids=[c.vector_id for c in chunks])
            except Exception as exc:
                logger.warning("Failed to delete vectors for doc %s: %s", doc.id, exc)
            for c in chunks:
                db.delete(c)
        db.delete(doc)
        db.commit()

    @staticmethod
    def process_document(db: Session, doc_id: int) -> dict:
        doc = db.get(KBDocument, doc_id)
        if not doc:
            raise ValueError(f"Document {doc_id} not found")
        doc.status = "processing"
        db.commit()
        try:
            text, _ = extract_text_from_file(doc.storage_path, doc.file_type)
            if not text.strip():
                doc.status = "failed"
                doc.error_message = "文件内容为空或无法提取文本"
                db.commit()
                return {"status": "failed", "message": "Empty or unreadable file"}

            kb = doc.kb
            chunks_data = chunk_text(text, kb.chunk_size, kb.chunk_overlap)
            if not chunks_data:
                doc.status = "failed"
                doc.error_message = "分块失败"
                db.commit()
                return {"status": "failed", "message": "No chunks produced"}

            embeddings = get_embeddings(kb.embedding_model)
            texts = [c["content"] for c in chunks_data]
            emb_list = embeddings.embed_documents(texts)

            folder = db.get(KBFolder, doc.folder_id) if doc.folder_id else None
            folder_path = KnowledgeBaseService.get_folder_path(db, folder) if folder else ""

            # Use the abstracted vector store backend
            vs = get_vector_store()
            collection_name = f"kb_{kb.id}"
            vs.ensure_collection(collection_name)

            chunk_records = []
            ids = []
            vectors = []
            documents = []
            metadatas = []
            for idx, (chunk_data, embedding) in enumerate(zip(chunks_data, emb_list)):
                vector_id = f"{doc.id}_chunk_{idx}_{uuid.uuid4().hex[:8]}"
                ids.append(vector_id)
                vectors.append(embedding)
                documents.append(chunk_data["content"])
                metadatas.append({
                    "document_id": str(doc.id),
                    "document_name": doc.original_filename,
                    "folder_path": folder_path,
                    "folder_id": str(doc.folder_id) if doc.folder_id else "",
                    "kb_id": str(kb.id),
                    "chunk_index": str(idx),
                    "page_number": "",
                    **(chunk_data.get("metadata") or {}),
                })
                chunk_records.append(KBChunk(
                    kb_id=kb.id, document_id=doc.id, folder_id=doc.folder_id,
                    vector_id=vector_id, chunk_index=idx,
                    content=chunk_data["content"],
                    metadata_=chunk_data.get("metadata") or {},
                ))

            # Batch upsert — much faster than per-chunk
            vs.upsert(collection_name, ids=ids, embeddings=vectors,
                      documents=documents, metadatas=metadatas)

            db.add_all(chunk_records)
            doc.status = "ready"
            doc.error_message = None
            db.commit()
            return {"status": "ready", "message": f"成功处理 {len(chunks_data)} 个分块", "chunks": len(chunks_data)}

        except Exception as exc:
            logger.exception("Failed to process document %d", doc_id)
            doc.status = "failed"
            doc.error_message = str(exc)
            db.commit()
            return {"status": "failed", "message": str(exc)}

    @staticmethod
    def search_knowledge_base(db: Session, kb_id: int, query: str,
                              top_k: int = 5, folder_id: int | None = None) -> list[dict]:
        kb = db.get(KnowledgeBase, kb_id)
        if not kb:
            raise ValueError(f"Knowledge base {kb_id} not found")

        embeddings = get_embeddings(kb.embedding_model)
        query_embedding = embeddings.embed_query(query)

        vs = get_vector_store()
        collection_name = f"kb_{kb.id}"

        where = {"kb_id": str(kb_id)}
        if folder_id is not None:
            folder = db.get(KBFolder, folder_id)
            if folder:
                folder_path = KnowledgeBaseService.get_folder_path(db, folder)
                where["folder_path"] = folder_path

        results = vs.query(
            collection_name,
            query_embeddings=[query_embedding],
            n_results=top_k,
            where=where,
        )

        hits = []
        ids_list = results.get("ids", [[]])
        distances_list = results.get("distances", [[]])
        metadatas_list = results.get("metadatas", [[]])
        for i in range(len(ids_list[0]) if ids_list else 0):
            vid = ids_list[0][i]
            distance = distances_list[0][i] if i < len(distances_list[0]) else 0
            metadata = metadatas_list[0][i] if i < len(metadatas_list[0]) else {}
            document_id = int(metadata.get("document_id", 0))
            doc = db.scalar(select(KBDocument).where(KBDocument.id == document_id))
            chunk = db.scalar(
                select(KBChunk).where(KBChunk.vector_id == vid, KBChunk.kb_id == kb_id)
            )
            hits.append({
                "vector_id": vid,
                "document_id": document_id,
                "document_name": doc.original_filename if doc else metadata.get("document_name", ""),
                "folder_path": metadata.get("folder_path", ""),
                "page_number": metadata.get("page_number"),
                "chunk_index": int(metadata.get("chunk_index", 0)) if metadata.get("chunk_index") else 0,
                "content": chunk.content if chunk else "",
                "score": round(1.0 - distance, 4) if distance is not None else 0.0,
            })
        return hits


# ============================================================
# User Management Service (Task 10)
# ============================================================

class UserService:
    @staticmethod
    def list_users(db: Session, user_id: int) -> list[User]:
        """Admin can list all users. Regular users can only see themselves."""
        admin_user = db.get(User, user_id)
        if admin_user and admin_user.is_superuser:
            return list(db.scalars(select(User).order_by(User.created_at)).all())
        return [db.get(User, user_id)]

    @staticmethod
    def update_user(db: Session, target_user: User, current_user_id: int, **kwargs) -> User:
        """Only admins can modify other users."""
        current = db.get(User, current_user_id)
        if not current.is_superuser and current.id != target_user.id:
            raise PermissionError("Only admins can modify other users.")
        for k, v in kwargs.items():
            if v is not None and hasattr(target_user, k):
                setattr(target_user, k, v)
        db.commit()
        db.refresh(target_user)
        return target_user

    @staticmethod
    def delete_user(db: Session, target_user: User, current_user_id: int) -> None:
        current = db.get(User, current_user_id)
        if not current.is_superuser:
            raise PermissionError("Only admins can delete users.")
        db.delete(target_user)
        db.commit()


# ============================================================
# System Settings Service (Task 11)
# ============================================================

class SystemSettingService:
    @staticmethod
    def get_setting(db: Session, key: str) -> SystemSetting | None:
        return db.scalar(select(SystemSetting).where(SystemSetting.key == key))

    @staticmethod
    def set_setting(db: Session, key: str, value: str, description: str = "") -> SystemSetting:
        setting = db.scalar(select(SystemSetting).where(SystemSetting.key == key))
        if setting:
            setting.value = value
            if description:
                setting.description = description
        else:
            setting = SystemSetting(key=key, value=value, description=description)
            db.add(setting)
        db.commit()
        db.refresh(setting)
        return setting

    @staticmethod
    def list_settings(db: Session) -> list[SystemSetting]:
        return list(db.scalars(select(SystemSetting).order_by(SystemSetting.key)).all())

    @staticmethod
    def delete_setting(db: Session, key: str) -> None:
        setting = db.scalar(select(SystemSetting).where(SystemSetting.key == key))
        if setting:
            db.delete(setting)
            db.commit()


# ============================================================
# Original utility functions (kept for backward compatibility)
# ============================================================

DEFAULT_SYSTEM_PROMPT = """You are a practical AI agent.
You can answer normally and call tools when they help.

AVAILABLE TOOLS:
1. calculator - Calculate numeric expressions
2. current_time - Get the current date/time for a timezone
3. list_workspace_files - List files in the project workspace
4. read_workspace_text_file - Read a text file from the workspace
5. search_knowledge_base - Search knowledge bases for relevant information

KNOWLEDGE BASE USAGE:
- When the user asks about company documents, policies, technical specs, or any stored knowledge, use search_knowledge_base.
- You can search all knowledge bases (omit kb_id) or a specific one (provide kb_id).
- Summarize the most relevant findings from the search results in your answer.
- If no knowledge base is available or the search returns nothing, let the user know.

TOOL PRIORITY (IMPORTANT):
- If a provided tool (including any MCP tool listed in the appended catalog) can directly answer the user's request, you MUST call that tool to obtain live/accurate data. Do not answer from prior knowledge when a tool is available and relevant.
- For factual, data, or real-time queries (e.g. train stations, weather, stock prices), always prefer the tool over guessing.

CHOICE INTERACTION:
- Only use `<blocks>` choice interaction when you genuinely need the USER to decide between options you cannot determine yourself.
- Never use it as a substitute for answering or for calling a tool, and never respond with empty content.
When you do use it, output your response text (never empty) followed by a <blocks> tag containing a JSON object with a choices array:
<blocks>{"choices": [{"label": "A. ???", "value": "A"}, {"label": "B. ???", "value": "B"}, {"label": "C. ???", "value": "C"}]}</blocks>

The text before <blocks> will be shown as regular message content. The choices will render as clickable buttons. When the user clicks a button, their selection (the value) will be sent as their message. Use labels that are clear and descriptive.
"""


def create_default_agent(db: Session, user_id: int) -> None:
    from app.models import AgentConfig
    from app.settings import get_settings
    settings = get_settings()
    agent = AgentConfig(
        user_id=user_id,
        name="Default Agent",
        description="Default local agent with time, calculator, workspace tools, and knowledge base search.",
        system_prompt=DEFAULT_SYSTEM_PROMPT,
        model_name=settings.openai_model,
        temperature=0,
        enabled=True,
    )
    db.add(agent)


def create_user(db: Session, email: str, username: str, password: str) -> User:
    user = User(email=email.lower(), username=username, password_hash=hash_password(password))
    db.add(user)
    db.flush()
    create_default_agent(db, user.id)
    # 注册即补基础角色权限（自动注册用户应看到基础菜单），否则权限驱动菜单下新用户几乎空白。
    from app.permissions import ensure_personal_defaults
    ensure_personal_defaults(user.id, db)
    # 自动授予所有默认角色（含 base 与后续新建的默认角色）——"设为默认角色→新用户自动获得"语义闭环。
    from app.rbac_seed import assign_default_roles_to_user
    assign_default_roles_to_user(db, user.id)
    db.commit()
    db.refresh(user)
    return user


def authenticate_user(db: Session, email: str, password: str) -> User | None:
    user = db.scalar(select(User).where(User.email == email.lower()))
    if not user or not verify_password(password, user.password_hash):
        return None
    return user


def new_thread_id() -> str:
    return f"thread-{uuid.uuid4().hex[:12]}"


# ============================================================
# RAG Core Services (RAG 澧炲己鏍稿績)
# ============================================================

class QueryRewriter:
    """Query rewriting for better retrieval."""

    @staticmethod
    def rewrite(query: str, kb_name: str = "") -> str:
        """Simple rule-based query rewriting."""
        stop_words = {"怎么", "如何", "什么", "为什么", "呢", "吧", "的", "人", "是", "在", "有", "我", "你", "他", "好", "它", "以", "还", "那", "中", "一", "上", "不", "都", "大", "得", "跟", "下", "对", "关于", "对于", "基于", "根据", "通过", "经过", "按照", "由于", "因为", "所以", "但是", "可是", "然而", "不过", "虽然", "尽管", "如果", "假如", "只要", "无论", "不管", "即使", "既然", "于是", "因此", "总之", "总而言之", "综上所述", "也就是说", "换句话说", "例如", "比如", "诸如", "像", "如同", "仿佛", "似的", "一样", "等等", "之类", "而言", "来说", "的话", "方面", "起来", "下来", "出来", "进来", "上去", "下去", "过来", "回去", "回来", "出去", "进去", "起来", "下来"}
        words = [w for w in query if w not in stop_words]
        return ''.join(words) if words else query


# ============================================================
# CrossEncoder cache — avoid reloading model on every retrieval
# ============================================================

_cross_encoder_cache: dict[str, Any] = {}


def _get_cross_encoder(model_name: str):
    """Get or create a cached CrossEncoder instance."""
    if model_name not in _cross_encoder_cache:
        try:
            from sentence_transformers import CrossEncoder
            _cross_encoder_cache[model_name] = CrossEncoder(model_name)
            logger.info("Loaded CrossEncoder model: %s", model_name)
        except Exception as exc:
            logger.warning("Failed to load CrossEncoder %s: %s", model_name, exc)
            return None
    return _cross_encoder_cache[model_name]


# ============================================================
# HybridRetriever — vector + keyword + RRF fusion + MMR + rerank
# ============================================================

class HybridRetriever:
    """Hybrid retrieval: vector + keyword + RRF fusion + MMR dedup + rerank."""

    def __init__(self, kb, db):
        self.kb = kb
        self.db = db
        self.vs = get_vector_store()
        # Use rag_config from the KB model, falling back to defaults
        self.rag_config = {**DEFAULT_RAG_CONFIG, **(getattr(kb, 'rag_config', None) or {})}

    def retrieve(self, query: str, top_k: int = 20, rerank_top_k: int = 10, folder_id=None):
        """Execute hybrid retrieval and return sorted results."""
        # Use config values if not explicitly overridden
        top_k = self.rag_config.get('top_k', top_k)
        rerank_top_k = self.rag_config.get('rerank_top_k', rerank_top_k)
        strategy = self.rag_config.get('retrieval_strategy', 'hybrid')

        vector_hits = []
        keyword_hits = []

        if strategy in ('hybrid', 'vector'):
            vector_hits = self._vector_search(query, top_k=top_k * 2, folder_id=folder_id)
        if strategy in ('hybrid', 'keyword'):
            keyword_hits = self._keyword_search(query, top_k=top_k * 2, folder_id=folder_id)

        if strategy == 'hybrid' and vector_hits and keyword_hits:
            fused = self._rrf_fusion(vector_hits, keyword_hits, k=self.rag_config.get('rrf_k', 60))
        elif vector_hits:
            fused = vector_hits
        elif keyword_hits:
            fused = keyword_hits
        else:
            return []

        # MMR deduplication
        if fused and self.rag_config.get('mmr_enabled', True):
            fused = self._mmr_deduplicate(fused, threshold=self.rag_config.get('mmr_threshold', 0.5))
        # Rerank
        if fused and self.rag_config.get('rerank_enabled', True):
            fused = self._rerank(query, fused[:rerank_top_k])
        # Filter low scores
        min_score = self.rag_config.get('min_relevance_score', 0.3)
        fused = [h for h in fused if h.get('score', 0) >= min_score]
        return fused[:top_k]

    def _vector_search(self, query, top_k, folder_id):
        """Vector search via the configured vector store backend."""
        try:
            embeddings = get_embeddings(self.kb.embedding_model)
            query_vec = embeddings.embed_query(query)
            collection_name = f"kb_{self.kb.id}"
            where = {"kb_id": str(self.kb.id)}
            if folder_id:
                where["folder_id"] = str(folder_id)
            results = self.vs.query(
                collection_name,
                query_embeddings=[query_vec],
                n_results=top_k,
                where=where,
            )
        except Exception as exc:
            logger.warning("Vector search failed: %s", exc)
            return []

        hits = []
        ids_list = results.get("ids", [[]])
        dist_list = results.get("distances", [[]])
        docs_list = results.get("documents", [[]])
        meta_list = results.get("metadatas", [[]])
        for i in range(len(ids_list[0]) if ids_list and ids_list[0] else 0):
            hits.append({
                'type': 'vector',
                'score': 1 - (dist_list[0][i] if i < len(dist_list[0]) else 0),
                'content': docs_list[0][i] if i < len(docs_list[0]) else '',
                'metadata': meta_list[0][i] if i < len(meta_list[0]) and meta_list[0][i] else {},
                'vector_id': ids_list[0][i],
            })
        return hits

    def _keyword_search(self, query, top_k, folder_id):
        """Keyword search via SQLite LIKE matching with proper tokenization."""
        # Fix: tokenize the query into words, not iterate over characters
        import re
        keywords = [w for w in re.split(r'[\s,，。.、；;！!？?]+', query) if len(w) > 1]
        if not keywords:
            keywords = [query]  # fallback to full query
        if not keywords:
            return []
        conditions = [KBChunk.content.like(f"%{kw}%") for kw in keywords]
        stmt = select(KBChunk).where(KBChunk.kb_id == self.kb.id, or_(*conditions)).limit(top_k)
        if folder_id:
            stmt = stmt.where(KBChunk.folder_id == folder_id)
        chunks = list(self.db.scalars(stmt).all())
        hits = []
        for chunk in chunks:
            doc = chunk.document
            score = sum(1 for kw in keywords if kw in chunk.content) / len(keywords)
            hits.append({
                'type': 'keyword',
                'score': score,
                'content': chunk.content,
                'metadata': {
                    'document_id': str(doc.id),
                    'document_name': doc.original_filename,
                    'folder_path': '',
                    'kb_id': str(self.kb.id),
                    'folder_id': str(chunk.folder_id) if chunk.folder_id else "",
                },
                'vector_id': chunk.vector_id,
            })
        return hits

    def _rrf_fusion(self, vector_hits, keyword_hits, k=60):
        """Reciprocal Rank Fusion."""
        rank_map = {}
        for i, hit in enumerate(vector_hits):
            vid = hit['vector_id']
            rank_map[vid] = rank_map.get(vid, 0) + k / (k + i + 1)
        for i, hit in enumerate(keyword_hits):
            vid = hit['vector_id']
            rank_map[vid] = rank_map.get(vid, 0) + k / (k + i + 1)
        merged = {}
        for hit in vector_hits + keyword_hits:
            vid = hit['vector_id']
            if vid not in merged:
                merged[vid] = {**hit, 'rrf_score': 0}
            merged[vid]['rrf_score'] = rank_map.get(vid, 0)
            if 'hit_source' not in merged[vid]:
                merged[vid]['hit_source'] = hit['type']
            elif hit['type'] != merged[vid]['hit_source']:
                merged[vid]['hit_source'] = 'both'
        return sorted(merged.values(), key=lambda x: x['rrf_score'], reverse=True)

    def _mmr_deduplicate(self, hits, threshold=0.5):
        """Maximal Marginal Relevance deduplication — optimized with batch embeddings."""
        if not hits:
            return []
        try:
            import numpy as np
            # Optimization: batch-embed all candidates at once instead of O(n²) individual calls
            embeddings = get_embeddings(self.kb.embedding_model)
            contents = [h['content'] for h in hits]
            all_embs = np.array(embeddings.embed_documents(contents))

            # Normalize embeddings for cosine similarity
            norms = np.linalg.norm(all_embs, axis=1, keepdims=True)
            norms[norms == 0] = 1
            normalized = all_embs / norms

            selected_indices = [0]
            remaining = list(range(1, len(hits)))
            while remaining and len(selected_indices) < 10:
                best_idx = 0
                best_score = -1
                for i, candidate_idx in enumerate(remaining):
                    # Vectorized similarity computation — no more per-pair API calls
                    sel_embs = normalized[selected_indices]
                    cand_emb = normalized[candidate_idx]
                    sims = sel_embs @ cand_emb
                    max_sim = float(sims.max())
                    mmr_score = hits[candidate_idx].get('score', 0) - threshold * max_sim
                    if mmr_score > best_score:
                        best_score = mmr_score
                        best_idx = i
                selected_indices.append(remaining.pop(best_idx))
            return [hits[i] for i in selected_indices]
        except Exception as exc:
            logger.warning("MMR dedup failed: %s", exc)
            return hits[:10]

    def _rerank(self, query, hits):
        """Cross-Encoder reranking with cached model."""
        if not hits:
            return []
        model_name = self.rag_config.get('rerank_model', 'BAAI/bge-reranker-base')
        ce = _get_cross_encoder(model_name)
        if ce is None:
            return sorted(hits, key=lambda x: x.get('rrf_score', 0), reverse=True)
        try:
            pairs = [[query, h['content']] for h in hits]
            scores = ce.predict(pairs)
            for hit, score in zip(hits, scores):
                hit['rerank_score'] = float(score)
                hit['score'] = float(score)
            return sorted(hits, key=lambda x: x['rerank_score'], reverse=True)
        except Exception as exc:
            logger.warning("Rerank failed: %s", exc)
            return sorted(hits, key=lambda x: x.get('rrf_score', 0), reverse=True)


class ContextBuilder:
    """Build LLM-ready context from retrieval results."""

    def __init__(self, max_tokens=4000):
        self.max_tokens = max_tokens

    def build(self, query, hits, include_sources=True):
        """Return (context_text, sources_list)."""
        if not hits:
            return "", []
        context_parts = []
        sources = []
        budget = self.max_tokens
        for i, hit in enumerate(hits, 1):
            content = hit['content']
            meta = hit.get('metadata', {})
            # Approximate token count
            token_count = len(content) // 3.5
            if token_count > budget:
                content = content[:int(budget * 3)] + '... [内容过长，已截断]'
                token_count = budget
            budget -= token_count
            if budget <= 0:
                break
            source_tag = ''
            if include_sources:
                doc_name = meta.get('document_name', 'Unknown')
                score_pct = hit.get('score', 0)
                folder_path = meta.get('folder_path', '')
                source_tag = f'[来源: {doc_name}, 相关度: {score_pct:.0%}]'
                if folder_path:
                    source_tag += f' ({folder_path})'
            context_parts.append(f"{source_tag}\n{content}\n")
            sources.append({
                'document_name': meta.get('document_name', ''),
                'folder_path': meta.get('folder_path', ''),
                'score': hit.get('score', 0),
                'rerank_score': hit.get('rerank_score'),
                'hit_source': hit.get('hit_source', 'vector'),
            })
        context = "\n=== 检索到的相关知识 ===\n\n" + "".join(context_parts)
        return context, sources


# RAG System Prompt
RAG_SYSTEM_PROMPT = """你是一个基于知识库的智能问答助手。

## 回答规则

1. **优先使用检索到的知识**: 当提供了检索结果时，必须基于这些内容回答问题
2. **必须引用来源**: 每个关键信息后面标注 [来源: 文件名]
3. **不知道就说不知道**: 如果检索结果中没有相关信息，明确告知用户
4. **不要编造答案**: 即使你觉得知道答案，也要以检索结果为准
5. **综合多来源**: 多个文档有相关信息时，综合后给出完整回答
6. **指出矛盾**: 不同文档有冲突信息时，告知用户并列出各方说法

## 回答风格

- 结构化，条理清晰
- 适当使用 Markdown 格式
- 引用具体数据和事实
- 如果问题超出知识库范围，告知用户并尝试用通用知识回答
"""


# ============================================================
# Provider Management Service
# ============================================================

class ProviderService:
    @staticmethod
    def create_provider(db: Session, user_id: int, name: str, base_url: str = "",
                        api_key: str = "", provider_type: str = "openai-compatible",
                        enabled: bool = True, is_default: bool = False) -> "Provider":
        from app.models import Provider
        if is_default:
            db.execute(
                sa_update(Provider).where(
                    Provider.user_id == user_id, Provider.is_default == True
                ).values(is_default=False)
            )
        provider = Provider(
            user_id=user_id, name=name, base_url=base_url, api_key=api_key,
            provider_type=provider_type, enabled=enabled, is_default=is_default,
        )
        db.add(provider)
        db.commit()
        db.refresh(provider)
        return provider

    @staticmethod
    def get_provider(db: Session, provider_id: int, user_id: int) -> "Provider | None":
        from app.models import Provider
        return db.scalar(
            select(Provider).where(Provider.id == provider_id, Provider.user_id == user_id)
        )

    @staticmethod
    def list_providers(db: Session, user_id: int) -> list["Provider"]:
        from app.models import Provider
        return list(db.scalars(
            select(Provider).where(Provider.user_id == user_id).order_by(Provider.created_at)
        ).all())

    @staticmethod
    def update_provider(db: Session, provider: "Provider", **kwargs) -> "Provider":
        for k, v in kwargs.items():
            if v is not None and hasattr(provider, k):
                if k == "is_default" and v:
                    db.execute(
                        sa_update(Provider).where(
                            Provider.user_id == provider.user_id,
                            Provider.id != provider.id,
                            Provider.is_default == True
                        ).values(is_default=False)
                    )
                setattr(provider, k, v)
        db.commit()
        db.refresh(provider)
        return provider

    @staticmethod
    def delete_provider(db: Session, provider: "Provider") -> None:
        db.delete(provider)
        db.commit()

    @staticmethod
    def create_model(db: Session, provider_id: int, model_name: str, model_type: str,
                     enabled: bool = True, is_default_chat: bool = False,
                     is_default_embedding: bool = False,
                     is_default_video: bool = False,
                     is_default_image: bool = False,
                     description: str = "") -> "ProviderModel":
        from app.models import ProviderModel
        # Check if model already exists — if so, update it (idempotent batch add)
        existing = db.scalar(
            select(ProviderModel).where(
                ProviderModel.provider_id == provider_id,
                ProviderModel.model_name == model_name
            )
        )
        if existing:
            existing.model_type = model_type
            existing.enabled = enabled
            existing.description = description or existing.description
            # Only flip default flags if the caller explicitly wants this model as default
            _default_flags = {
                "chat": "is_default_chat",
                "embedding": "is_default_embedding",
                "video": "is_default_video",
                "image": "is_default_image",
            }
            if any(getattr(existing, v) for v in _default_flags.values() if v in _default_flags):
                pass  # keep existing defaults if already set
            elif is_default_chat or is_default_embedding or is_default_video or is_default_image:
                # Apply caller's default flags only if none are already set
                existing.is_default_chat = is_default_chat
                existing.is_default_embedding = is_default_embedding
                existing.is_default_video = is_default_video
                existing.is_default_image = is_default_image
            db.commit()
            db.refresh(existing)
            return existing
        # Reset other defaults of the same type
        _default_flags_map = {
            "chat": ("is_default_chat", is_default_chat),
            "embedding": ("is_default_embedding", is_default_embedding),
            "video": ("is_default_video", is_default_video),
            "image": ("is_default_image", is_default_image),
        }
        if model_type in _default_flags_map:
            flag_name, flag_value = _default_flags_map[model_type]
            if flag_value:
                db.execute(
                    sa_update(ProviderModel).where(
                        ProviderModel.provider_id == provider_id,
                        ProviderModel.model_type == model_type,
                    ).values({flag_name: False})
                )
        pm = ProviderModel(
            provider_id=provider_id, model_name=model_name, model_type=model_type,
            enabled=enabled, is_default_chat=is_default_chat,
            is_default_embedding=is_default_embedding,
            is_default_video=is_default_video,
            is_default_image=is_default_image,
            description=description,
        )
        db.add(pm)
        db.commit()
        db.refresh(pm)
        return pm

    @staticmethod
    def get_provider_models(db: Session, provider_id: int) -> list["ProviderModel"]:
        from app.models import ProviderModel
        return list(db.scalars(
            select(ProviderModel).where(
                ProviderModel.provider_id == provider_id
            ).order_by(ProviderModel.model_type, ProviderModel.model_name)
        ).all())

    @staticmethod
    def update_model(db: Session, model: "ProviderModel", **kwargs) -> "ProviderModel":
        for k, v in kwargs.items():
            if v is not None and hasattr(model, k):
                setattr(model, k, v)
        db.commit()
        db.refresh(model)
        return model

    @staticmethod
    def delete_model(db: Session, model: "ProviderModel") -> None:
        db.delete(model)
        db.commit()

    @staticmethod
    def get_default_chat_model(db: Session, provider_id: int) -> str | None:
        from app.models import ProviderModel
        return db.scalar(
            select(ProviderModel.model_name).where(
                ProviderModel.provider_id == provider_id,
                ProviderModel.model_type == "chat",
                ProviderModel.is_default_chat == True,
                ProviderModel.enabled == True,
            )
        )

    @staticmethod
    def get_default_embedding_model(db: Session, provider_id: int) -> str | None:
        from app.models import ProviderModel
        return db.scalar(
            select(ProviderModel.model_name).where(
                ProviderModel.provider_id == provider_id,
                ProviderModel.model_type == "embedding",
                ProviderModel.is_default_embedding == True,
                ProviderModel.enabled == True,
            )
        )

    @staticmethod
    def get_default_model_by_type(db: Session, provider_id: int, model_type: str) -> str | None:
        """Get the default model name for a specific model type from a provider."""
        from app.models import ProviderModel
        type_to_flag = {
            "chat": ProviderModel.is_default_chat,
            "image": ProviderModel.is_default_image,
            "video": ProviderModel.is_default_video,
            "embedding": ProviderModel.is_default_embedding,
        }
        flag = type_to_flag.get(model_type)
        if flag is None:
            return None
        return db.scalar(
            select(ProviderModel.model_name).where(
                ProviderModel.provider_id == provider_id,
                ProviderModel.model_type == model_type,
                flag == True,
                ProviderModel.enabled == True,
            )
        )

    @staticmethod
    def get_default_model(db: Session, user_id: int) -> "DefaultModelResponse":
        from app.models import Provider, ProviderModel
        from app.schemas import DefaultModelResponse
        default_provider = db.scalar(
            select(Provider).where(Provider.user_id == user_id, Provider.is_default == True)
        )
        if not default_provider:
            return DefaultModelResponse(
                chat_model=None, embedding_model=None,
                video_model=None, image_model=None,
                provider_id=None, provider_name=None,
            )
        chat_model = ProviderService.get_default_chat_model(db, default_provider.id)
        embedding_model = ProviderService.get_default_embedding_model(db, default_provider.id)
        video_model = ProviderService.get_default_model_by_type(db, default_provider.id, "video")
        image_model = ProviderService.get_default_model_by_type(db, default_provider.id, "image")
        return DefaultModelResponse(
            chat_model=chat_model,
            embedding_model=embedding_model,
            video_model=video_model,
            image_model=image_model,
            provider_id=default_provider.id,
            provider_name=default_provider.name,
        )

